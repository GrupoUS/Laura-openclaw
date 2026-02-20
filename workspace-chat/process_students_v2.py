import os
import pandas as pd
from docx import Document
import re
import unicodedata
import shutil

# Configurações
BASE_DIR = "/Users/mauricio/.openclaw/alunos"

# Mapeamento fornecido (Curso -> Arquivo -> Turma/Ciclo)
MAPS = {
    "33": {
        "Mentoria e Pós TRINTAE3 (Turma 3).xlsx": "Turma 3",
        "Mentoria e Pós TRINTAE3 (Turma 4).xlsx": "Turma 4",
        "Mentoria e Pós TRINTAE3 (Turma 5).xlsx": "Turma 5",
        "Mentoria e Pós TRINTAE3 (Turma 6).xlsx": "Turma 6",
    },
    "NEON": {
        "NEON (Ciclo 1).xlsx": "Ciclo 1",
        "NEON (Ciclo 2).xlsx": "Ciclo 2",
        "NEON (Ciclo 3).xlsx": "Ciclo 3",
    },
    "OTB": {
        "OTB - Out of The Box.xlsx": "Ciclo Unico",
    }
}

def slugify(text):
    if not text or pd.isna(text):
        return "sem_nome"
    text = str(text)
    text = unicodedata.normalize('NFD', text).encode('ascii', 'ignore').decode('utf-8')
    text = re.sub(r'[^\w\s-]', '', text).strip()
    return text

def create_docx(row, output_path):
    doc = Document()
    doc.add_heading('Ficha do Aluno', 0)
    for col in row.index:
        val = row[col]
        if pd.isna(val): val = ""
        doc.add_paragraph(f"{col}: {val}")
    doc.save(output_path)

def main():
    # Limpeza preventiva (opcional, mas seguro para garantir a nova estrutura)
    # Aqui vamos apenas processar e sobrescrever/criar.
    
    for curso, files in MAPS.items():
        for filename, turma in files.items():
            source_path = os.path.join(BASE_DIR, filename)
            if not os.path.exists(source_path):
                print(f"⚠️ Arquivo não encontrado: {filename}")
                continue

            print(f"📖 Processando {curso} -> {turma} ({filename})...")
            try:
                df = pd.read_excel(source_path)
            except Exception as e:
                print(f"❌ Erro ao ler {filename}: {e}")
                continue

            name_col = next((c for c in df.columns if 'nome' in c.lower()), df.columns[0])
            
            for _, row in df.iterrows():
                aluno_nome = slugify(row[name_col])
                # Estrutura: alunos/{CURSO}/{TURMA_OU_CICLO}/{NOME_DO_ALUNO}/
                target_dir = os.path.join(BASE_DIR, curso, turma, aluno_nome)
                os.makedirs(target_dir, exist_ok=True)
                
                docx_path = os.path.join(target_dir, f"{aluno_nome}.docx")
                create_docx(row, docx_path)
                print(f"  ✅ Gerado: {curso}/{turma}/{aluno_nome}")

if __name__ == "__main__":
    main()
