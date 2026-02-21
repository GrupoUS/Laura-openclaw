import os
import json
import requests
import pandas as pd
from docx import Document
from pathlib import Path
import re
import shutil

# API Configs
ASAAS_KEY = "$aact_prod_000MzkwODA2MWY2OGM3MWRlMDU2NWM3MzJlNzZmNGZhZGY6OjE0ZjlmNDVjLWQ3NzItNDU1ZS1hNGRlLWRkZDUyMTM2OWNjMTo6JGFhY2hfZjEyNDcxYzctNTM4OC00ZjlkLWEwOWMtNDY1NDMwMmQ3ODA5"
ASAAS_URL = "https://api.asaas.com/v3"

KIWIFY_CONFIG = {
    "client_id": "cbd7d830-fc76-45ae-836c-a1662a04ad78",
    "client_secret": "def66f463b9963de81f099626a3222892dcb513e6bade2ae5fedee924a4c3f93",
    "account_id": "hRsfEMMkkwZELXF",
    "base_url": "https://public-api.kiwify.com"
}

BASE_DIR = "/Users/mauricio/.openclaw/alunos"

def get_kiwify_token():
    url = f"{KIWIFY_CONFIG['base_url']}/v1/oauth/token"
    data = {
        "client_id": KIWIFY_CONFIG["client_id"],
        "client_secret": KIWIFY_CONFIG["client_secret"]
    }
    try:
        r = requests.post(url, data=data)
        return r.json().get("access_token")
    except:
        return None

KIWIFY_TOKEN = get_kiwify_token()

def asaas_get_payments(email=None, cpf=None):
    headers = {"access_token": ASAAS_KEY}
    params = {}
    if cpf:
        params["cpfCnpj"] = re.sub(r"\D", "", str(cpf))
    elif email:
        params["email"] = email
    else:
        return []

    try:
        r = requests.get(f"{ASAAS_URL}/customers", headers=headers, params=params)
        customers = r.json().get("data", [])
        if not customers:
            return []
        
        customer_id = customers[0]["id"]
        r_pay = requests.get(f"{ASAAS_URL}/payments", headers=headers, params={"customer": customer_id, "limit": 100})
        return r_pay.json().get("data", [])
    except:
        return []

def kiwify_search(query):
    if not KIWIFY_TOKEN: return []
    url = f"{KIWIFY_CONFIG['base_url']}/v1/sales"
    headers = {
        "Authorization": f"Bearer {KIWIFY_TOKEN}",
        "x-kiwify-account-id": KIWIFY_CONFIG["account_id"]
    }
    try:
        r = requests.get(url, headers=headers, params={"page_size": 100})
        if r.status_code != 200: return []
        sales = r.json().get("data", [])
        found = []
        q = str(query).lower()
        for s in sales:
            c = s.get("customer", {})
            if q in (c.get("email") or "").lower() or q in (c.get("mobile") or ""):
                found.append(s)
        return found
    except:
        return []

def get_student_info_from_docx(path):
    try:
        doc = Document(path)
        info = {}
        for para in doc.paragraphs:
            if ":" in para.text:
                parts = para.text.split(":", 1)
                if len(parts) == 2:
                    info[parts[0].strip().lower()] = parts[1].strip()
        return info
    except:
        return {}

def enrich_student(student_dir, course, turma):
    docx_files = list(Path(student_dir).glob("*.docx"))
    if not docx_files: return
    
    main_docx = docx_files[0]
    info = get_student_info_from_docx(main_docx)
    
    email = info.get("email") or info.get("e-mail")
    cpf = info.get("cpf") or info.get("documento")
    phone = info.get("telefone") or info.get("whatsapp")
    
    print(f"  🔍 Enriquecendo: {student_dir.split('/')[-1]} ({email})")
    
    asaas_payments = asaas_get_payments(email, cpf)
    kiwify_sales = kiwify_search(email or phone or main_docx.stem)
    
    # Financial Summary
    fin_path = os.path.join(student_dir, "Ficha_Financeira.docx")
    doc = Document()
    doc.add_heading(f"Resumo Financeiro - {main_docx.stem}", 0)
    
    found_turma_2 = False

    if asaas_payments:
        doc.add_heading("Asaas", level=1)
        # Parcela logic
        received = [p for p in asaas_payments if p["status"] in ["RECEIVED", "CONFIRMED"]]
        pending = [p for p in asaas_payments if p["status"] in ["PENDING", "OVERDUE"]]
        
        doc.add_paragraph(f"Parcelas Pagas: {len(received)}")
        doc.add_paragraph(f"Parcelas Pendentes: {len(pending)}")
        doc.add_paragraph(f"Total Pago: R$ {sum(p['value'] for p in received):.2f}")
        doc.add_paragraph(f"Total a Pagar: R$ {sum(p['value'] for p in pending):.2f}")
        
        for p in asaas_payments:
            desc = p.get("description", "") or ""
            if "TURMA 2" in desc.upper():
                found_turma_2 = True
            doc.add_paragraph(f"• {p.get('dueDate')}: R$ {p.get('value')} - {p.get('status')} ({desc})")

    if kiwify_sales:
        doc.add_heading("Kiwify", level=1)
        for s in kiwify_sales:
            doc.add_paragraph(f"Produto: {s.get('product', {}).get('name')}")
            doc.add_paragraph(f"Status: {s.get('status')}")
            doc.add_paragraph(f"Valor: R$ {s.get('net_amount') / 100 if s.get('net_amount') else 0:.2f}")
            doc.add_paragraph("-" * 20)

    doc.save(fin_path)
    
    # OTB Turma Logic
    if course == "OTB" and turma == "Ciclo Unico":
        new_turma = "Turma 2" if found_turma_2 else "Turma 1"
        new_dir = os.path.join(BASE_DIR, "OTB", new_turma, os.path.basename(student_dir))
        os.makedirs(os.path.dirname(new_dir), exist_ok=True)
        if student_dir != new_dir:
            print(f"    ➔ Movendo OTB para {new_turma}")
            shutil.move(student_dir, new_dir)

def main():
    for course_dir in os.listdir(BASE_DIR):
        cp = os.path.join(BASE_DIR, course_dir)
        if not os.path.isdir(cp): continue
        
        for turma_dir in os.listdir(cp):
            tp = os.path.join(cp, turma_dir)
            if not os.path.isdir(tp): continue
            
            print(f"📂 Turma: {course_dir}/{turma_dir}")
            # Use list() to avoid issues when moving directories
            for student_folder in list(os.listdir(tp)):
                sp = os.path.join(tp, student_folder)
                if not os.path.isdir(sp): continue
                enrich_student(sp, course_dir, turma_dir)

if __name__ == "__main__":
    main()
