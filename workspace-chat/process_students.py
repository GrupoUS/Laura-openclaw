import os
import pandas as pd
from docx import Document
import re
import unicodedata
from pathlib import Path

# Configurações
BASE_SOURCE_DIR = "/Users/mauricio/.openclaw/alunos"
BASE_TARGET_DIR = "/Users/mauricio/.openclaw/alunos"

# Mapeamento fornecido
FILE_MAPPING = {
    "Mentoria e Pós TRINTAE3 (Turma 3).xlsx": "33/Turma 3",
    "Mentoria e Pós TRINTAE3 (Turma 4).xlsx": "33/Turma 4",
    "Mentoria e Pós TRINTAE3 (Turma 5).xlsx": "33/Turma 5",
    "Mentoria e Pós TRINTAE3 (Turma 6).xlsx": "33/Turma 6",
    "NEON (Ciclo 1).xlsx": "NEON/Ciclo 1",
    "NEON (Ciclo 2).xlsx": "NEON/Ciclo 2",
    "NEON (Ciclo 3).xlsx": "NEON/Ciclo 3",
    "OTB - Out of The Box.xlsx": "OTB"
}

def slugify(text):
    if not text or pd.isna(text):
        return "sem_nome"
    text = str(text)
    # Remove acentos
    text = unicodedata.normalize('NFD', text).encode('ascii', 'ignore').decode('utf-8')
    # Remove caracteres especiais
    text = re.sub(r'[^\w\s-]', '', text).strip()
    return text

def create_docx(row, output_path):
    doc = Document()
    doc.add_heading('Ficha do Aluno', 0)
    
    for col in row.index:
        val = row[col]
        if pd.isna(val):
            val = ""
        doc.add_paragraph(f"{col}: {val}")
    
    doc.save(output_path)

def process_file(filename, relative_path):
    source_path = os.path.join(BASE_SOURCE_DIR, filename)
    if not os.path.exists(source_path):
        print(f"⚠️ Arquivo não encontrado: {source_path}")
        return

    print(f"📖 Processando {filename}...")
    try:
        df = pd.read_excel(source_path)
    except Exception as e:
        print(f"❌ Erro ao ler {filename}: {e}")
        return

    # Tenta encontrar a coluna de nome (pode variar entre planilhas)
    name_col = next((c for c in df.columns if 'nome' in c.lower()), df.columns[0])
    
    target_dir = os.path.join(BASE_TARGET_DIR, relative_path)
    
    for _, row in df.iterrows():
        aluno_nome = slugify(row[name_col])
        aluno_dir = os.path.join(target_dir, aluno_nome)
        
        os.makedirs(aluno_dir, exist_ok=True)
        
        docx_path = os.path.join(aluno_dir, f"{aluno_nome}.docx")
        create_docx(row, docx_path)
        print(f"  ✅ Gerado: {relative_path}/{aluno_nome}")

def main():
    for filename, rel_path in FILE_MAPPING.items():
        process_file(filename, rel_path)

if __name__ == "__main__":
    main()
