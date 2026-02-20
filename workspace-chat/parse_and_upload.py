import os
import pandas as pd
from docx import Document
import requests
import json
import base64
import re
import unicodedata

# Configurações
GITHUB_TOKEN = "ghp_YyQ7F0z_F_V_Z_C_4_F_V_Z_C_4_F_V_Z_C" # placeholder, need to find the real one
GITHUB_REPO = "GrupoUS/Laura-openclaw"
BRANCH = "main"

# Mapeamento fornecido
FILE_MAPPING = {
    "Mentoria e Pós TRINTAE3 (Turma 3).xlsx": "alunos/33/Turma 3/",
    "Mentoria e Pós TRINTAE3 (Turma 4).xlsx": "alunos/33/Turma 4/",
    "Mentoria e Pós TRINTAE3 (Turma 5).xlsx": "alunos/33/Turma 5/",
    "Mentoria e Pós TRINTAE3 (Turma 6).xlsx": "alunos/33/Turma 6/",
    "NEON (Ciclo 1).xlsx": "alunos/NEON/Ciclo 1/",
    "NEON (Ciclo 2).xlsx": "alunos/NEON/Ciclo 2/",
    "NEON (Ciclo 3).xlsx": "alunos/NEON/Ciclo 3/",
    "OTB - Out of The Box.xlsx": "alunos/OTB/"
}

def slugify(text):
    if not text or pd.isna(text):
        return "sem_nome"
    text = str(text)
    text = unicodedata.normalize('NFD', text).encode('ascii', 'ignore').decode('utf-8')
    text = re.sub(r'[^\w\s-]', '', text).strip().lower()
    return re.sub(r'[-\s]+', '_', text)

def create_docx(row, output_path):
    doc = Document()
    doc.add_heading('Ficha do Aluno', 0)
    
    for col in row.index:
        val = row[col]
        if pd.isna(val):
            val = ""
        doc.add_paragraph(f"{col}: {val}")
    
    doc.save(output_path)

def main():
    # TODO: Implement GitHub API interaction
    pass

if __name__ == "__main__":
    main()
