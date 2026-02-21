import os
import json
import requests
import pandas as pd
from docx import Document
from pathlib import Path
import re

# API Configs
ASAAS_KEY = "$aact_prod_000MzkwODA2MWY2OGM3MWRlMDU2NWM3MzJlNzZmNGZhZGY6OjE0ZjlmNDVjLWQ3NzItNDU1ZS1hNGRlLWRkZDUyMTM2OWNjMTo6JGFhY2hfZjEyNDcxYzctNTM4OC00ZjlkLWEwOWMtNDY1NDMwMmQ3ODA5"
ASAAS_URL = "https://api.asaas.com/v3"

KIWIFY_CONFIG = {
    "client_id": "cbd7d830-fc76-45ae-836c-a1662a04ad78",
    "client_secret": "def66f463b9963de81f099626a3222892dcb513e6bade2ae5fedee924a4c3f93",
    "account_id": "hRsfEMMkkwZELXF",
    "base_url": "https://public-api.kiwify.com"
}

BASE_DIR = "/Users/mauricio/.openclaw/alunos"

def get_kiwify_token():
    url = f"{KIWIFY_CONFIG['base_url']}/v1/oauth/token"
    data = {
        "client_id": KIWIFY_CONFIG["client_id"],
        "client_secret": KIWIFY_CONFIG["client_secret"]
    }
    r = requests.post(url, data=data)
    return r.json().get("access_token")

KIWIFY_TOKEN = get_kiwify_token()

def asaas_get_payments(email=None, cpf=None):
    headers = {"access_token": ASAAS_KEY}
    params = {}
    if cpf:
        params["cpfCnpj"] = re.sub(r"\D", "", cpf)
    elif email:
        params["email"] = email
    else:
        return []

    r = requests.get(f"{ASAAS_URL}/customers", headers=headers, params=params)
    customers = r.json().get("data", [])
    if not customers:
        return []
    
    customer_id = customers[0]["id"]
    r_pay = requests.get(f"{ASAAS_URL}/payments", headers=headers, params={"customer": customer_id, "limit": 100})
    return r_pay.json().get("data", [])

def kiwify_search(query):
    url = f"{KIWIFY_CONFIG['base_url']}/v1/sales"
    headers = {
        "Authorization": f"Bearer {KIWIFY_TOKEN}",
        "x-kiwify-account-id": KIWIFY_CONFIG["account_id"]
    }
    try:
        r = requests.get(url, headers=headers, params={"page_size": 100})
        if r.status_code != 200:
            print(f"  ⚠️ Kiwify API Error ({r.status_code})")
            return []
        sales = r.json().get("data", [])
    except Exception as e:
        print(f"  ⚠️ Kiwify Search Failed: {e}")
        return []
    
    found = []
    q = str(query).lower()
    for s in sales:
        c = s.get("customer", {})
        if q in (c.get("email") or "").lower() or q in (c.get("mobile") or ""):
            found.append(s)
    return found

def get_student_info_from_docx(path):
    try:
        doc = Document(path)
        info = {}
        for para in doc.paragraphs:
            if ":" in para.text:
                key, val = para.text.split(":", 1)
                info[key.strip().lower()] = val.strip()
        return info
    except:
        return {}

def enrich_student(student_dir, course, turma):
    # Find the docx file
    docx_files = list(Path(student_dir).glob("*.docx"))
    if not docx_files:
        return
    
    main_docx = docx_files[0]
    info = get_student_info_from_docx(main_docx)
    
    email = info.get("email") or info.get("e-mail")
    cpf = info.get("cpf") or info.get("documento")
    phone = info.get("telefone") or info.get("whatsapp")
    
    print(f"  🔍 Buscando dados para: {main_docx.stem} ({email})")
    
    # Asaas Data
    asaas_payments = asaas_get_payments(email, cpf)
    # Kiwify Data
    kiwify_sales = kiwify_search(email or phone or main_docx.stem)
    
    if not asaas_payments and not kiwify_sales:
        return

    # Create Financial Doc
    fin_path = os.path.join(student_dir, "Ficha_Financeira.docx")
    doc = Document()
    doc.add_heading(f"Resumo Financeiro - {main_docx.stem}", 0)
    
    if asaas_payments:
        doc.add_heading("Asaas", level=1)
        total_paid = sum(p["value"] for p in asaas_payments if p["status"] == "RECEIVED" or p["status"] == "CONFIRMED")
        total_pending = sum(p["value"] for p in asaas_payments if p["status"] == "PENDING" or p["status"] == "OVERDUE")
        
        doc.add_paragraph(f"Total Pago: R$ {total_paid:.2f}")
        doc.add_paragraph(f"Total Pendente/Atrasado: R$ {total_pending:.2f}")
        
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vencimento'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Descrição'
        
        for p in asaas_payments:
            row_cells = table.add_row().cells
            row_cells[0].text = p.get("dueDate", "")
            row_cells[1].text = str(p.get("value", ""))
            row_cells[2].text = p.get("status", "")
            row_cells[3].text = p.get("description", "") or ""

    if kiwify_sales:
        doc.add_heading("Kiwify", level=1)
        for s in kiwify_sales:
            doc.add_paragraph(f"Produto: {s.get('product', {}).get('name')}")
            doc.add_paragraph(f"Status: {s.get('status')}")
            doc.add_paragraph(f"Valor: R$ {s.get('net_amount') / 100 if s.get('net_amount') else 0:.2f}")
            doc.add_paragraph(f"Data: {s.get('created_at')}")
            doc.add_paragraph("-" * 20)

    doc.save(fin_path)
    print(f"  ✅ Ficha financeira criada: {fin_path}")

def main():
    for course_dir in os.listdir(BASE_DIR):
        cp = os.path.join(BASE_DIR, course_dir)
        if not os.path.isdir(cp): continue
        
        for turma_dir in os.listdir(cp):
            tp = os.path.join(cp, turma_dir)
            if not os.path.isdir(tp): continue
            
            print(f"📂 Processando Turma: {course_dir} / {turma_dir}")
            for student_dir in os.listdir(tp):
                sp = os.path.join(tp, student_dir)
                if not os.path.isdir(sp): continue
                enrich_student(sp, course_dir, turma_dir)

if __name__ == "__main__":
    main()
