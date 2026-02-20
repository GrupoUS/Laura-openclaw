"""Content extraction for DOCX files."""

import io
from docx import Document
from docx.oxml.ns import qn

from src.utils.logging import get_logger

logger = get_logger(__name__)


class DocxExtractor:
    """Extract text content from DOCX files."""

    def extract(self, content: bytes) -> str:
        """
        Extract text from DOCX bytes.

        Returns Markdown-formatted text preserving headings.
        """
        try:
            doc = Document(io.BytesIO(content))

            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Convert heading styles to Markdown
                if para.style.name.startswith("Heading"):
                    level = self._get_heading_level(para.style.name)
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)

            # Extract tables
            for table in doc.tables:
                parts.append(self._table_to_markdown(table))

            result = "\n\n".join(parts)
            logger.debug("Extracted DOCX", paragraphs=len(doc.paragraphs), chars=len(result))
            return result

        except Exception as e:
            logger.error("DOCX extraction failed", error=str(e))
            raise

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            # "Heading 1" -> 1, "Heading 2" -> 2, etc.
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 1

    def _table_to_markdown(self, table) -> str:
        """Convert a DOCX table to Markdown format."""
        rows = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                rows.append(separator)

        return "\n".join(rows)
